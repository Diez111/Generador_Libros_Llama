import re
import requests
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# Token y URL base de la API de Llama
api_token = "998e75d9-108c-4da5-82fb-a1ee0e9f9f05"
base_url = "https://api.llama-api.com/"

def llama_request(messages, model="llama3-8b", stream=False):
    """Realiza la petición a la API de Llama."""
    data = {
        "messages": messages,
        "model": model,
        "stream": stream
    }
    headers = {
        "Authorization": f"Bearer {api_token}",
        "Content-Type": "application/json"
    }
    response = requests.post(base_url + "chat/completions", json=data, headers=headers)
    if response.status_code == 200:
        return response.json()
    else:
        print(f"Error: {response.status_code}")
        print(response.text)
        return None

def generate_story_blueprint(theme):
    """
    Genera el blueprint en formato JSON a partir del tema.
    Este blueprint contiene detalles de personajes, eventos y el mundo narrativo, y se utiliza para mantener la coherencia.
    """
    messages = [
        {
            "role": "system",
            "content": (
                "Eres un narrador profesional. Crea un blueprint en formato JSON que incluya una descripción detallada de los personajes principales, "
                "eventos clave y el mundo en el que se desarrolla la historia. Registra todos los eventos importantes para que la IA pueda consultarlos y mantener la coherencia. "
                "No incluyas el blueprint en la narrativa final y no debes repetir frases o oraciones al inicio, debes cuidar de no ser repetitivo con las palabras o oraciones que utilizas."
            )
        },
        {
            "role": "user",
            "content": f"Genera un blueprint en JSON para una historia cuyo tema es: {theme}"
        }
    ]
    result = llama_request(messages)
    if result and "choices" in result and result["choices"]:
        return result["choices"][0]["message"]["content"]
    return "{}"

def clean_text(text):
    """
    Realiza el post-procesamiento del texto:
      - Elimina encabezados tipo "Capítulo X" y marcas de formato.
      - Elimina líneas vacías duplicadas.
    """
    text = re.sub(r'(?im)^\s*Cap[ií]tulo\s+\d+(:\s*.*)?\s*$', '', text)
    text = text.replace("**", "")
    text = re.sub(r'\n\s*\n', '\n\n', text)
    return text.strip()

def deduplicate_text(accumulated, new_segment):
    """
    Compara el nuevo segmento con el contenido acumulado y elimina frases repetidas.
    Esta función usa una comparación simple basada en frases.
    """
    # Dividimos el nuevo segmento en oraciones y filtramos aquellas ya presentes en el acumulado
    sentences = re.split(r'(?<=[.!?])\s+', new_segment)
    filtered = [s for s in sentences if s.strip() not in accumulated]
    return " ".join(filtered)

def get_word_count(text):
    """Retorna el número de palabras en el texto."""
    return len(text.split())

def generate_chapter_segment(chapter_number, total_chapters, blueprint, word_goal, initial=True):
    """
    Genera un segmento del capítulo. Si es el segmento inicial se incluyen instrucciones completas;
    de lo contrario, se pide continuar la narrativa sin reiniciar ideas o insertar encabezados.
    """
    if initial:
        base_instruction = (
            f"Genera el capítulo {chapter_number} de la historia utilizando el siguiente blueprint. "
            f"El capítulo debe tener AL MENOS {word_goal} palabras en total. "
            "No repitas ideas ni insertes encabezados o listas de eventos en la narrativa, no debes repetir oraciones ni palabras, debes usar un vocabulario variado. "
            "Mantén un tono uniforme y evita redundancias."
        )
        if chapter_number in (total_chapters, total_chapters - 1):
            base_instruction += (
                " La historia se aproxima a su final; concluye la trama de forma coherente sin reiniciar la narrativa."
            )
        user_content = f"{base_instruction}\n\nBlueprint actual: {blueprint}"
    else:
        user_content = (
            "Continúa la narrativa exactamente donde se dejó, sin reiniciar ideas ni insertar nuevos encabezados o metadatos. "
            "Evita repetir lo ya mencionado y añade nuevos detalles para alcanzar el mínimo de palabras requerido."
        )
    messages = [
        {
            "role": "system",
            "content": (
                "Eres un escritor profesional y creativo. Genera contenido narrativo de alta calidad que se integre de forma fluida con lo ya generado. "
                "No incluyas encabezados, etiquetas o metadatos en el cuerpo del texto y evita la repetición innecesaria."
            )
        },
        {
            "role": "user",
            "content": user_content
        }
    ]
    result = llama_request(messages)
    if result and "choices" in result and result["choices"]:
        segment = result["choices"][0]["message"]["content"]
        return clean_text(segment)
    return ""

def generate_full_chapter(chapter_number, total_chapters, blueprint, chapter_length):
    """
    Genera el capítulo completo de manera iterativa hasta alcanzar el mínimo de palabras.
    Se integra el contenido de forma fluida y se evita la repetición utilizando deduplicación.
    """
    if chapter_length in ["1", "2"]:
        word_goal = 2100
    elif chapter_length == "3":
        word_goal = 6000
    else:
        word_goal = 2100

    chapter = generate_chapter_segment(chapter_number, total_chapters, blueprint, word_goal, initial=True)
    accumulated = chapter
    iterations = 0
    max_iterations = 10

    while get_word_count(accumulated) < word_goal and iterations < max_iterations:
        print(f"Extensión del capítulo {chapter_number}, iteración {iterations+1} (palabras: {get_word_count(accumulated)})")
        continuation = generate_chapter_segment(chapter_number, total_chapters, blueprint, word_goal, initial=False)
        # Deduplicar para evitar repeticiones
        filtered = deduplicate_text(accumulated, continuation)
        if not filtered or get_word_count(filtered) < 50:  # Si la continuación es mínima o ya repetida
            break
        accumulated += "\n\n" + filtered
        iterations += 1

    final_chapter = clean_text(accumulated)
    print(f"Capítulo {chapter_number} finalizado con {get_word_count(final_chapter)} palabras.")
    return final_chapter

def update_blueprint(current_blueprint, chapter_text, chapter_number):
    """
    Actualiza el blueprint en formato JSON incorporando los eventos y cambios narrativos del capítulo.
    Solo devuelve el blueprint actualizado, sin insertar el contenido en la narrativa.
    """
    messages = [
        {
            "role": "system",
            "content": (
                "Eres un planificador de historias extremadamente detallado. A partir del blueprint actual y del contenido del capítulo, "
                "actualiza el blueprint en formato JSON para reflejar todos los eventos importantes, cambios en personajes y detalles relevantes. "
                "No incluyas el texto narrativo en el blueprint, solo la información estructurada para consulta interna."
            )
        },
        {
            "role": "user",
            "content": (
                f"Blueprint actual: {current_blueprint}\n\n"
                f"Capítulo {chapter_number} generado: {chapter_text}\n\n"
                "Actualiza el blueprint en formato JSON para mantener la coherencia narrativa y registrar los eventos importantes."
            )
        }
    ]
    result = llama_request(messages)
    if result and "choices" in result and result["choices"]:
        return result["choices"][0]["message"]["content"]
    return current_blueprint

def main():
    title = input("Ingresa el título del libro: ")
    author = input("Ingresa el autor del libro: ")
    theme = input("Ingresa el tema del libro: ")

    print("Selecciona la longitud de los capítulos:")
    print("1 - Corto (mínimo 2100 palabras)")
    print("2 - Medio (mínimo 2100 palabras)")
    print("3 - Largo (mínimo 6000 palabras)")
    chapter_length = input("Ingresa 1, 2 o 3: ").strip()

    try:
        num_chapters = int(input("Ingresa el número total de capítulos a generar: "))
    except ValueError:
        print("Número de capítulos inválido.")
        return

    # Generar blueprint inicial
    blueprint = generate_story_blueprint(theme)
    if not blueprint:
        print("No se pudo generar el blueprint de la historia.")
        return
    print("Blueprint inicial generado.")

    # Crear documento Word
    doc = Document()

    # Página de Título
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.font.size = Pt(28)
    title_run.bold = True
    doc.add_page_break()

    # Página de Autor
    author_paragraph = doc.add_paragraph()
    author_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_run = author_paragraph.add_run(author)
    author_run.font.size = Pt(22)
    author_run.italic = True
    doc.add_page_break()

    # Índice
    index_paragraph = doc.add_paragraph("ÍNDICE\n")
    for i in range(1, num_chapters+1):
        index_paragraph.add_run(f"Capítulo {i}\n")
    doc.add_page_break()

    # Generar cada capítulo y actualizar el blueprint
    current_blueprint = blueprint
    for i in range(1, num_chapters+1):
        print(f"\n--- Generando Capítulo {i} ---")
        chapter_text = generate_full_chapter(i, num_chapters, current_blueprint, chapter_length)
        # Agregar encabezado del capítulo (solo en el documento final)
        chapter_heading = doc.add_heading(f"Capítulo {i}", level=1)
        chapter_heading.runs[0].bold = True
        doc.add_paragraph(chapter_text)
        doc.add_page_break()
        current_blueprint = update_blueprint(current_blueprint, chapter_text, i)

    # Guardar documento en carpeta actual
    output_filename = title.replace(" ", "_") + ".docx"
    doc.save(output_filename)
    print(f"\nLibro generado exitosamente: {output_filename}")

if __name__ == "__main__":
    main()
